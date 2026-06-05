"""Extract content from Word documents."""

import os
from typing import Dict, Any, List
from docx import Document


def extract_docx(filepath: str) -> Dict[str, Any]:
    """
    Extract all content from a DOCX file.
    Returns: {
        'paragraphs': [str],
        'tables': [[[str]]],
        'sections': [{heading: str, content: str}],
        'filename': str
    }
    """
    doc = Document(filepath)

    paragraphs = []
    tables = []
    sections = []
    current_section = {'heading': '', 'content': []}

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Paragraph element
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            paragraphs.append(text)

            if para.style.name.startswith('Heading'):
                if current_section['heading'] or current_section['content']:
                    sections.append({
                        'heading': current_section['heading'],
                        'content': '\n'.join(current_section['content'])
                    })
                current_section = {'heading': text, 'content': []}
            else:
                current_section['content'].append(text)

        elif tag == 'tbl':
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table is not None:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append(rows)

    # Don't forget last section
    if current_section['heading'] or current_section['content']:
        sections.append({
            'heading': current_section['heading'],
            'content': '\n'.join(current_section['content'])
        })

    return {
        'paragraphs': paragraphs,
        'tables': tables,
        'sections': sections,
        'filename': os.path.basename(filepath),
    }


def build_text_summary(data: Dict[str, Any]) -> str:
    """Build a readable text summary from extracted DOCX data."""
    lines = [f"# {data['filename']}\n"]

    for section in data['sections']:
        if section['heading']:
            lines.append(f"\n## {section['heading']}")
        if section['content']:
            lines.append(section['content'])

    for i, table in enumerate(data.get('tables', [])):
        if not table:
            continue
        lines.append(f"\n### 表格 {i+1}")
        if table[0]:
            lines.append("\n| " + " | ".join(table[0]) + " |")
            lines.append("|" + "|".join(["---"] * len(table[0])) + "|")
        for row in table[1:]:
            cols = len(table[0]) if table else len(row)
            padded = row + [''] * (cols - len(row))
            lines.append("| " + " | ".join(padded[:cols]) + " |")

    return '\n'.join(lines)
